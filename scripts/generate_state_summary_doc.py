from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


REPORT_DATE_TEXT = "February 20, 2026"
REPORT_DATE_FILE = "2026-02-20"


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def add_number(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Number")


def build_doc() -> Path:
    out_dir = Path("docs")
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / f"VoicePractice_State_Summary_{REPORT_DATE_FILE}.docx"

    doc = Document()

    title = doc.add_paragraph()
    title_run = title.add_run("VoicePracticeApp - State Summary and Platform/Process Review (Updated)")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    date_p = doc.add_paragraph()
    date_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_p.add_run(f"Date: {REPORT_DATE_TEXT}").italic = True
    doc.add_paragraph("")

    add_heading(doc, "1. Executive Snapshot", 1)
    doc.add_paragraph(
        "The product is now very close to the intended pre-live shape. Core app behavior is stable, admin control "
        "surfaces are broad and functional, and the platform has moved to a cleaner hosted architecture: API and DB on Render, "
        "admin web on Vercel, source/deploy control through GitHub, and Android app delivery through Expo/EAS."
    )
    doc.add_paragraph(
        "Most recent high-impact work focused on operational reliability and environment clarity rather than feature sprawl. "
        "The biggest architecture change since the prior summary is the database move away from Supabase connection usage to "
        "Render Postgres as the active backing store for the API."
    )

    add_heading(doc, "2. Vision and Product Direction (Current Understanding)", 1)
    add_heading(doc, "2.1 Near-Term (Now to Initial Live)", 2)
    add_bullet(doc, "Stable and trustworthy app behavior is the highest priority.")
    add_bullet(doc, "Admin Utility remains the source of truth for managed app content and scenarios.")
    add_bullet(doc, "Feature additions are secondary to reliability, polish, and deployment discipline.")
    add_bullet(doc, "Scenario breadth should scale through data/content, not major code rewrites.")

    add_heading(doc, "2.2 Mid-Term", 2)
    add_bullet(doc, "Selective integrations (billing/bookkeeping export and simulation analytics export) are likely.")
    add_bullet(doc, "Capacity growth should support moderate concurrency (40-50 concurrent users target).")
    add_bullet(doc, "Architecture should remain intentionally narrow, not become a large multi-product platform.")

    add_heading(doc, "3. Current Platform Flow (As Deployed/Used)", 1)
    add_heading(doc, "3.1 GitHub", 2)
    add_bullet(doc, "Canonical source of truth for code, history, and deployment trigger commits.")
    add_bullet(doc, "Main branch pushes drive Render and Vercel redeploy behavior.")
    add_bullet(doc, "CI/verifications exist for fast and broad checks (verify:fast and verify:all).")

    add_heading(doc, "3.2 Render", 2)
    add_bullet(doc, "Hosts the API service: https://voicepractice-api-dev.onrender.com")
    add_bullet(doc, "Now also hosts the active Postgres database used by the API.")
    add_bullet(doc, "API readiness/health checks are used operationally (/ready and /health).")
    add_bullet(doc, "Render env controls runtime behavior (DATABASE_URL, STORAGE_PROVIDER, pool/timeouts, secrets).")

    add_heading(doc, "3.3 Vercel", 2)
    add_bullet(doc, "Hosts Admin Utility frontend.")
    add_bullet(doc, "Frontend points to Render API base URL via env config.")

    add_heading(doc, "3.4 Expo / EAS (Android)", 2)
    add_bullet(doc, "Managed Expo workflow remains in place.")
    add_bullet(doc, "APK builds are performed via EAS from VS Code terminal command flow.")
    add_bullet(doc, "Backend-only changes generally do not require a new APK build.")

    add_heading(doc, "3.5 Supabase (Current Status)", 2)
    add_bullet(doc, "Supabase is no longer the active DB path for this environment.")
    add_bullet(doc, "Supabase org remains available temporarily for rollback comfort, but not in active env wiring.")
    add_bullet(doc, "Current direction is decommission after short stabilization period and final confidence checks.")

    add_heading(doc, "4. Detailed Change Log Since Previous Summary", 1)
    add_heading(doc, "4.1 Reliability Troubleshooting and DB Connectivity", 2)
    add_bullet(doc, "Investigated intermittent 'Database temporarily unavailable' events in admin/API paths.")
    add_bullet(doc, "Observed readiness and config check failures tied to connection timeout/certificate chain issues.")
    add_bullet(doc, "Validated endpoint behavior repeatedly with terminal curl checks.")
    add_bullet(doc, "Concluded that active hosted DB path needed to be simplified and stabilized.")

    add_heading(doc, "4.2 Database Platform Move: Supabase-Referenced Connection -> Render Postgres", 2)
    add_bullet(doc, "Provisioned Render Postgres instance in the same region footprint as API service.")
    add_bullet(doc, "Shifted API DATABASE_URL to Render Postgres.")
    add_bullet(doc, "Kept STORAGE_PROVIDER=postgres so persistence remains durable in hosted DB.")
    add_bullet(doc, "Redeployed API and validated health/readiness operational status.")
    add_bullet(doc, "Upgraded DB from free to Basic-256mb tier for more stable testing/runtime behavior.")

    add_heading(doc, "4.3 Data Hygiene and Seed Removal (Admin Utility)", 2)
    add_bullet(doc, "Removed runtime demo reseeding behavior from API normalization path.")
    add_bullet(doc, "Added purge path for previously inserted demo rows (demo orgs/users/sessions/scores/support cases/tokens).")
    add_bullet(doc, "Removed default starter org seed and purge any existing org_starter record.")
    add_bullet(doc, "Result: clean admin data surface with no forced dummy accounts.")

    add_heading(doc, "4.4 Security/Operations Follow-up", 2)
    add_bullet(doc, "OpenAI API key rotation recommended after key visibility in screenshot context.")
    add_bullet(doc, "API key replacement process clarified to avoid downtime (create new, deploy, verify, then revoke old).")
    add_bullet(doc, "Uptime monitoring confirms current health/readiness endpoints are green.")

    add_heading(doc, "4.5 Final Stabilization Updates (Today)", 2)
    add_bullet(doc, "Fixed mobile onboarding initialization loop risk by gating app bootstrap to one initial pass.")
    add_bullet(doc, "Added admin enterprise account join-request moderation UI (approve, approve-as-org-admin, reject).")
    add_bullet(doc, "Added personal-account user delete action in admin Personal tab with double confirmation.")
    add_bullet(doc, "Fixed API /users/:userId delete restriction so personal users can be deleted by platform admin.")
    add_bullet(doc, "Hardened mobile update stream resync behavior after API restarts/cursor reset conditions.")
    add_bullet(doc, "Added explicit app auto-forward from Org Access screen to Home when org approval converts account to enterprise.")
    add_bullet(doc, "Observed live behavior now aligns with expected flow: onboarding -> verify -> request -> approve -> enterprise activation.")

    add_heading(doc, "5. Supabase to Render Move - Deep Technical/Operational Summary", 1)
    doc.add_paragraph(
        "This section captures the migration in practical terms, why it happened, what changed, and what it means going forward."
    )

    add_heading(doc, "5.1 Why the Move Was Made", 2)
    add_bullet(
        doc,
        "Operational goal: remove ongoing connectivity instability and reduce troubleshooting complexity during pre-live phase."
    )
    add_bullet(
        doc,
        "Observed symptom pattern: intermittent DB unavailability despite successful deploys, plus repeated timeout/certificate-related errors."
    )
    add_bullet(
        doc,
        "Risk posture: for a live-ready posture, stable and predictable DB connectivity is more important than provider diversification."
    )

    add_heading(doc, "5.2 What Changed (Exactly)", 2)
    add_bullet(doc, "Database provider path changed to Render Postgres for active environment.")
    add_bullet(doc, "Render API env DATABASE_URL now references Render-managed Postgres instance.")
    add_bullet(doc, "No mobile API endpoint change required; app continues to call same Render API host.")
    add_bullet(doc, "No schema migration framework introduced; app continues using current app_state JSONB table model.")

    add_heading(doc, "5.3 What Did Not Change", 2)
    add_bullet(doc, "API endpoint contract remained stable for mobile and admin frontend.")
    add_bullet(doc, "Admin Utility functionality and content model remained intact.")
    add_bullet(doc, "Expo/EAS Android build path remained unchanged.")
    add_bullet(doc, "Vercel hosting model remained unchanged.")

    add_heading(doc, "5.4 Reliability/Performance Implications", 2)
    add_bullet(doc, "Fewer moving parts in active hosted path (API + DB both in Render estate).")
    add_bullet(doc, "Reduced ambiguity when diagnosing DB outages (single provider control plane).")
    add_bullet(doc, "Basic paid DB tier provides a more dependable baseline than ephemeral free-db behavior.")
    add_bullet(doc, "Still requires normal production disciplines: backups, restore drills, and staged change control.")

    add_heading(doc, "5.5 Cost and Scaling Implications", 2)
    add_bullet(doc, "Current DB spend is intentionally low (basic entry tier) for pre-live and early alpha.")
    add_bullet(doc, "Scaling path is linear: increase DB/service tier as user load and concurrent sessions grow.")
    add_bullet(doc, "This tier is suitable for current testing/demo and small early cohorts; monitor usage and latency before expansion.")

    add_heading(doc, "5.6 Operational Runbook (Now)", 2)
    add_number(doc, "Push code to GitHub main.")
    add_number(doc, "Render API redeploys from latest commit.")
    add_number(doc, "Render DB persists data across API redeploys.")
    add_number(doc, "Validate /ready and /health after deploy.")
    add_number(doc, "Hard refresh admin utility and run quick smoke flow (login, content/accounts checks).")
    add_number(doc, "Only rebuild APK when mobile code/build config changes, not for backend-only changes.")

    add_heading(doc, "6. Current Functional State (Admin + Mobile + API)", 1)
    add_heading(doc, "6.1 Content/Admin Model", 2)
    add_bullet(doc, "Industry/Role/Role Industries/Scenarios model is active.")
    add_bullet(doc, "Searchable selectors and role-industry activation behavior are in place.")
    add_bullet(doc, "Scenario table actions and CSV export patterns are implemented.")
    add_bullet(doc, "Delete confirmations include stronger intent confirmation patterns where configured.")

    add_heading(doc, "6.2 Accounts and Billing Surfaces", 2)
    add_bullet(doc, "Accounts tab includes comprehensive CSV export support.")
    add_bullet(doc, "Enterprise account detail includes Usage/Billing surface with contract detail management.")
    add_bullet(doc, "Soft limit concepts and usage visualization paths are part of current model.")

    add_heading(doc, "6.3 Support/Logs", 2)
    add_bullet(doc, "Support screen supports sorting by headers and date-range CSV export.")
    add_bullet(doc, "Auto-created support/error reporting flows were added previously and remain active.")

    add_heading(doc, "6.4 Session and Security Behavior", 2)
    add_bullet(doc, "Org-level max simulation length support is present.")
    add_bullet(doc, "Session duration visibility and red warning when near limit were implemented.")
    add_bullet(doc, "Admin token/session handling was hardened with sign-out and inactivity behavior.")

    add_heading(doc, "6.5 Enterprise Join Lifecycle Behavior", 2)
    add_bullet(doc, "Join requests now surface directly on enterprise account detail pages in admin utility.")
    add_bullet(doc, "Platform admin can approve as normal enterprise user or directly as org admin.")
    add_bullet(doc, "After approval, mobile app now transitions out of join-request holding state without manual app restart.")
    add_bullet(doc, "Mobile long-poll update stream now recovers reliably from backend in-memory cursor resets.")

    add_heading(doc, "6.6 Validation Status", 2)
    add_bullet(doc, "Automated critical-flow test passed locally end-to-end (admin login, org create, onboard, verify, join request, approval, user delete, audit checks).")
    add_bullet(doc, "Fast verification suite continues to pass (shared/api/admin build + mobile type-check).")
    add_bullet(doc, "Live endpoint probes remain healthy when checked (/health, /ready).")

    add_heading(doc, "7. Known Risks and How They Are Being Managed", 1)
    add_heading(doc, "7.1 Mobile minimatch advisory", 2)
    add_bullet(doc, "Still tracked as transitive tooling dependency risk in Expo/RN dependency tree.")
    add_bullet(doc, "Current assessment: tooling-chain risk, not direct runtime application logic exploit path.")
    add_bullet(doc, "Mitigation strategy: keep upstream updates monitored and avoid destabilizing platform shift pre-launch.")

    add_heading(doc, "7.2 Remaining pre-live hardening work", 2)
    add_bullet(doc, "Formal staging/prod split remains recommended when budget/process are ready.")
    add_bullet(doc, "Backup/restore drill should be documented and practiced.")
    add_bullet(doc, "Load test should be run before larger user onboarding.")
    add_bullet(doc, "Email provider integration still pending for true production communication flows.")

    add_heading(doc, "8. How Collaboration/Execution Is Currently Working", 1)
    add_number(doc, "User runs account-authenticated commands (provider portals, EAS login/build, sensitive env updates).")
    add_number(doc, "Assistant handles code updates, verification, commit/push, and exact command guidance.")
    add_number(doc, "User shares terminal outputs/screens when needed; assistant performs iterative troubleshooting.")
    add_number(doc, "Final state checks are based on endpoint probes, deploy status, and in-app verification.")

    add_heading(doc, "9. Immediate Next Steps (Practical and Ordered)", 1)
    add_number(doc, "Complete current redeploy verification and confirm clean admin data state.")
    add_number(doc, "Keep Supabase disabled from env use; decommission after short confidence window.")
    add_number(doc, "Run short soak checks on /ready and /health after each major config change.")
    add_number(doc, "Capture and store an environment matrix (Render API, Render DB, Vercel, EAS project, key owners).")
    add_number(doc, "Plan staging model and load test sequence before broader alpha onboarding.")

    add_heading(doc, "10. Appendix - Recent Relevant Commits", 1)
    add_bullet(doc, "b5b1a40 - harden mobile updates resync + auto-forward after org approval")
    add_bullet(doc, "2b2c6b9 - allow admin deletion of personal users in API")
    add_bullet(doc, "7e1bfdb - add org join moderation UI + personal delete action in admin-web")
    add_bullet(doc, "3418d0c - prevent mobile app re-initialization loop during onboarding")
    add_bullet(doc, "9a261e3 - remove demo seed data and purge existing demo records")
    add_bullet(doc, "e40274d - remove starter enterprise seed from default data")
    add_bullet(doc, "1a8a777 - remove direct accounts shortcut from admin landing")
    add_bullet(doc, "099bec7 - CI release gates and dependency automation baseline")

    add_heading(doc, "11. Quick Command Reference", 1)
    doc.add_paragraph("Repository root validation commands:")
    doc.add_paragraph("npm run build:api")
    doc.add_paragraph("npm run verify:fast")
    doc.add_paragraph("npm run verify:all")
    doc.add_paragraph("curl -sS https://voicepractice-api-dev.onrender.com/ready")
    doc.add_paragraph("curl -sS -i https://voicepractice-api-dev.onrender.com/health")

    add_heading(doc, "12. Handoff Prompt for Next Chat", 1)
    doc.add_paragraph(
        "Use docs/VoicePractice_State_Summary_2026-02-20.docx as current source context. "
        "Architecture is Render API + Render Postgres + Vercel admin + Expo/EAS Android. "
        "Admin Utility content model is source of truth. Focus on launch-stability-first changes, avoid unnecessary platform shifts."
    )

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    path = build_doc()
    print(path)
